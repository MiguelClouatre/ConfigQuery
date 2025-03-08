"""
Document Processor Module

Handles the ingestion of documents into ChromaDB, including:
- Loading documents from files (TXT, PDF, DOCX, XLSX/XLS)
- Chunking text into manageable pieces
- Extracting metadata
- Generating embeddings
- Storing in ChromaDB
"""
import os
import time
from pathlib import Path
from datetime import datetime
import uuid
import chromadb
from sentence_transformers import SentenceTransformer
import config

# Import libraries for handling different file types
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX processing will not be available.")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not installed. PDF processing will not be available.")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("Warning: pandas not installed. Excel processing will not be available.")

# Initialize the embedding model - reuse your existing model
model = SentenceTransformer('all-MiniLM-L6-v2')

# Initialize ChromaDB client - reuse your existing connection
client = chromadb.PersistentClient(path=config.CHROMA_DB_PATH)
collection = client.get_or_create_collection(config.CHROMA_COLLECTION_NAME)

class Document:
    """
    Represents a document with text content and metadata.
    """
    def __init__(self, text, metadata=None):
        """
        Initialize a document with text and metadata.
        
        Args:
            text (str): The full text content of the document
            metadata (dict, optional): Metadata about the document. Defaults to None.
        """
        self.text = text
        self.metadata = metadata or {}
        self.chunks = []
        self.chunk_metadata = []
    
    def add_metadata(self, key, value):
        """Add a metadata field to the document."""
        self.metadata[key] = value

def load_text_file(file_path):
    """
    Load a text file and create a Document object.
    
    Args:
        file_path (str): Path to the text file
        
    Returns:
        Document: A Document object with the file's text and basic metadata
    """
    path = Path(file_path)
    
    # Check if file exists
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        # Read the file
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Create document with basic metadata
        doc = Document(text)
        
        # Add file metadata
        doc.add_metadata('filename', path.name)
        doc.add_metadata('file_path', str(path))
        doc.add_metadata('file_type', 'text')
        doc.add_metadata('file_size_bytes', path.stat().st_size)
        
        # Add creation and modification dates if available
        doc.add_metadata('creation_date', datetime.fromtimestamp(path.stat().st_ctime).isoformat())
        doc.add_metadata('modification_date', datetime.fromtimestamp(path.stat().st_mtime).isoformat())
        
        # Add ingestion timestamp
        doc.add_metadata('ingestion_date', datetime.now().isoformat())
        
        return doc
    
    except UnicodeDecodeError:
        raise ValueError(f"File is not encoded in UTF-8: {file_path}")
    except Exception as e:
        raise Exception(f"Error loading file {file_path}: {str(e)}")

def load_pdf_file(file_path):
    """
    Load a PDF file and create a Document object.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        Document: A Document object with the file's text and basic metadata
    """
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 is not installed. Cannot process PDF files.")
    
    path = Path(file_path)
    
    # Check if file exists
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        # Open the PDF file
        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Extract text from all pages
            text = ""
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n\n"
            
            # Create document with basic metadata
            doc = Document(text)
            
            # Add file metadata
            doc.add_metadata('filename', path.name)
            doc.add_metadata('file_path', str(path))
            doc.add_metadata('file_type', 'pdf')
            doc.add_metadata('file_size_bytes', path.stat().st_size)
            doc.add_metadata('page_count', len(reader.pages))
            
            # Add creation and modification dates if available
            doc.add_metadata('creation_date', datetime.fromtimestamp(path.stat().st_ctime).isoformat())
            doc.add_metadata('modification_date', datetime.fromtimestamp(path.stat().st_mtime).isoformat())
            
            # Add ingestion timestamp
            doc.add_metadata('ingestion_date', datetime.now().isoformat())
            
            return doc
    
    except Exception as e:
        raise Exception(f"Error loading PDF file {file_path}: {str(e)}")

def load_docx_file(file_path):
    """
    Load a DOCX file and create a Document object.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        Document: A Document object with the file's text and basic metadata
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Cannot process DOCX files.")
    
    path = Path(file_path)
    
    # Check if file exists
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        # Open the DOCX file
        doc_obj = docx.Document(path)
        
        # Extract text from all paragraphs
        text = "\n\n".join([para.text for para in doc_obj.paragraphs if para.text.strip()])
        
        # Create document with basic metadata
        doc = Document(text)
        
        # Add file metadata
        doc.add_metadata('filename', path.name)
        doc.add_metadata('file_path', str(path))
        doc.add_metadata('file_type', 'docx')
        doc.add_metadata('file_size_bytes', path.stat().st_size)
        
        # Add creation and modification dates if available
        doc.add_metadata('creation_date', datetime.fromtimestamp(path.stat().st_ctime).isoformat())
        doc.add_metadata('modification_date', datetime.fromtimestamp(path.stat().st_mtime).isoformat())
        
        # Add ingestion timestamp
        doc.add_metadata('ingestion_date', datetime.now().isoformat())
        
        return doc
    
    except Exception as e:
        raise Exception(f"Error loading DOCX file {file_path}: {str(e)}")

def load_excel_file(file_path):
    """
    Load an Excel file (XLSX, XLS) and create a Document object.
    
    Args:
        file_path (str): Path to the Excel file
        
    Returns:
        Document: A Document object with the file's text and basic metadata
    """
    if not PANDAS_AVAILABLE:
        raise ImportError("pandas is not installed. Cannot process Excel files.")
    
    path = Path(file_path)
    
    # Check if file exists
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(path)
        sheet_names = excel_file.sheet_names
        
        # Process each sheet
        all_text_parts = []
        
        for sheet_name in sheet_names:
            # Read the sheet into a dataframe
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            
            # Add sheet header
            all_text_parts.append(f"Sheet: {sheet_name}")
            
            # Get column headers
            headers = df.columns.tolist()
            header_text = "Columns: " + ", ".join([str(h) for h in headers])
            all_text_parts.append(header_text)
            
            # Convert dataframe to text, preserving table structure
            rows_text = []
            for index, row in df.iterrows():
                # Format each cell value and join with tabs
                row_values = [f"{col}: {str(val)}" for col, val in row.items()]
                rows_text.append(" | ".join(row_values))
            
            # Join rows with newlines
            if rows_text:
                table_text = "\n".join(rows_text)
                all_text_parts.append(table_text)
            else:
                all_text_parts.append("(Empty sheet)")
            
            # Add a separator between sheets
            all_text_parts.append("-" * 50)
        
        # Join all text parts with double newlines
        text = "\n\n".join(all_text_parts)
        
        # Create document with basic metadata
        doc = Document(text)
        
        # Add file metadata
        doc.add_metadata('filename', path.name)
        doc.add_metadata('file_path', str(path))
        doc.add_metadata('file_type', 'excel')
        doc.add_metadata('file_size_bytes', path.stat().st_size)
        doc.add_metadata('sheet_count', len(sheet_names))
        doc.add_metadata('sheet_names', ", ".join(sheet_names))
        
        # Add creation and modification dates if available
        doc.add_metadata('creation_date', datetime.fromtimestamp(path.stat().st_ctime).isoformat())
        doc.add_metadata('modification_date', datetime.fromtimestamp(path.stat().st_mtime).isoformat())
        
        # Add ingestion timestamp
        doc.add_metadata('ingestion_date', datetime.now().isoformat())
        
        return doc
    
    except Exception as e:
        raise Exception(f"Error loading Excel file {file_path}: {str(e)}")

def chunk_text(document, chunk_size=1000, chunk_overlap=200):
    """
    Split document text into chunks of approximately equal size with overlap.
    
    Args:
        document (Document): The document to chunk
        chunk_size (int, optional): Target size of each chunk in characters. Defaults to 1000.
        chunk_overlap (int, optional): Overlap between chunks in characters. Defaults to 200.
        
    Returns:
        Document: The document with chunks added
    """
    text = document.text
    
    # If the text is empty or whitespace only, create a single empty chunk
    if not text or text.isspace():
        document.chunks = ["[Empty document]"]
        chunk_metadata_entry = document.metadata.copy()
        chunk_metadata_entry['chunk_index'] = 0
        chunk_metadata_entry['chunk_id'] = str(uuid.uuid4())
        document.chunk_metadata = [chunk_metadata_entry]
        return document
    
    # If the text is shorter than chunk_size, keep it as one chunk
    if len(text) <= chunk_size:
        document.chunks = [text]
        
        # Create metadata for single chunk
        chunk_metadata_entry = document.metadata.copy()
        chunk_metadata_entry['chunk_index'] = 0
        chunk_metadata_entry['chunk_id'] = str(uuid.uuid4())
        
        document.chunk_metadata = [chunk_metadata_entry]
        return document
    
    chunks = []
    chunk_metadata = []
    
    # Split text into paragraphs
    paragraphs = text.split('\n\n')
    
    # Handle case with no paragraph breaks
    if len(paragraphs) <= 1:
        # Just split by a fixed size
        for i in range(0, len(text), chunk_size - chunk_overlap):
            chunk = text[i:i + chunk_size]
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
                
                # Create metadata for this chunk
                chunk_metadata_entry = document.metadata.copy()
                chunk_metadata_entry['chunk_index'] = len(chunks) - 1
                chunk_metadata_entry['chunk_id'] = str(uuid.uuid4())
                chunk_metadata.append(chunk_metadata_entry)
                
        document.chunks = chunks
        document.chunk_metadata = chunk_metadata
        return document
    
    current_chunk = []
    current_size = 0
    
    for paragraph in paragraphs:
        # Clean up the paragraph
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        
        # If adding this paragraph would exceed chunk_size, save the current chunk and start a new one
        paragraph_size = len(paragraph)
        
        if current_size + paragraph_size > chunk_size and current_chunk:
            # Save the current chunk
            chunks.append('\n\n'.join(current_chunk))
            
            # Create metadata for this chunk
            chunk_metadata_entry = document.metadata.copy()
            chunk_metadata_entry['chunk_index'] = len(chunks) - 1
            chunk_metadata_entry['chunk_id'] = str(uuid.uuid4())
            chunk_metadata.append(chunk_metadata_entry)
            
            # Start a new chunk with overlap
            # Include the last few paragraphs from the previous chunk for context
            overlap_size = 0
            overlap_paragraphs = []
            
            # Go backwards through the current chunk to find paragraphs for overlap
            for p in reversed(current_chunk):
                if overlap_size + len(p) <= chunk_overlap:
                    overlap_paragraphs.insert(0, p)
                    overlap_size += len(p)
                else:
                    break
            
            current_chunk = overlap_paragraphs
            current_size = overlap_size
        
        # Add the paragraph to the current chunk
        current_chunk.append(paragraph)
        current_size += paragraph_size
    
    # Don't forget the last chunk
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
        
        # Create metadata for the last chunk
        chunk_metadata_entry = document.metadata.copy()
        chunk_metadata_entry['chunk_index'] = len(chunks) - 1
        chunk_metadata_entry['chunk_id'] = str(uuid.uuid4())
        chunk_metadata.append(chunk_metadata_entry)
    
    document.chunks = chunks
    document.chunk_metadata = chunk_metadata
    
    return document

def store_document_in_chroma(document):
    """
    Store a document's chunks and metadata in ChromaDB.
    
    Args:
        document (Document): The document to store
        
    Returns:
        dict: Results info with number of chunks stored
    """
    if not document.chunks:
        raise ValueError("Document has no chunks. Call chunk_text() first.")
    
    if not document.chunk_metadata:
        raise ValueError("Document has no chunk metadata. There may be an issue with chunking.")
    
    # Generate embeddings for each chunk
    embeddings = model.encode(document.chunks).tolist()
    
    # Prepare IDs, documents, embeddings, and metadata for ChromaDB
    ids = [f"{document.metadata.get('filename', 'doc')}_{meta['chunk_id']}" for meta in document.chunk_metadata]
    
    # Add the chunks to ChromaDB
    collection.add(
        ids=ids,
        embeddings=embeddings,
        documents=document.chunks,
        metadatas=document.chunk_metadata
    )
    
    return {
        "chunks_added": len(document.chunks),
        "document_name": document.metadata.get('filename', 'Unnamed document'),
        "document_size": document.metadata.get('file_size_bytes', 0)
    }

def process_file(file_path, chunk_size=1000, chunk_overlap=200):
    """
    Process a file and add it to ChromaDB based on its type.
    
    Args:
        file_path (str): Path to the file
        chunk_size (int, optional): Target size of each chunk in characters. Defaults to 1000.
        chunk_overlap (int, optional): Overlap between chunks in characters. Defaults to 200.
        
    Returns:
        dict: Results information
    """
    try:
        path = Path(file_path)
        file_extension = path.suffix.lower()
        
        # Load document based on file type
        if file_extension == '.txt':
            document = load_text_file(file_path)
        elif file_extension == '.pdf' and PDF_AVAILABLE:
            document = load_pdf_file(file_path)
        elif file_extension in ['.docx', '.doc'] and DOCX_AVAILABLE:
            document = load_docx_file(file_path)
        elif file_extension in ['.xlsx', '.xls'] and PANDAS_AVAILABLE:
            document = load_excel_file(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_extension}"
            }
        
        # Chunk the document
        document = chunk_text(document, chunk_size, chunk_overlap)
        
        # Store in ChromaDB
        store_document_in_chroma(document)
        
        # Return success info
        return {
            "success": True,
            "filename": document.metadata.get('filename', 'Unnamed document'),
            "chunks": len(document.chunks),
            "size_kb": document.metadata.get('file_size_bytes', 0) / 1024,
            "file_type": document.metadata.get('file_type', 'unknown')
        }
    
    except Exception as e:
        # Return error info
        return {
            "success": False,
            "error": str(e)
        }

# For backward compatibility
def process_text_file(file_path, chunk_size=1000, chunk_overlap=200):
    """
    Process a single text file and add it to ChromaDB.
    (Maintained for backward compatibility)
    
    Args:
        file_path (str): Path to the text file
        chunk_size (int, optional): Target size of each chunk in characters. Defaults to 1000.
        chunk_overlap (int, optional): Overlap between chunks in characters. Defaults to 200.
        
    Returns:
        dict: Results information
    """
    return process_file(file_path, chunk_size, chunk_overlap)

# Install dependencies function to help users
def install_dependencies():
    """
    Install required dependencies for full document processing capabilities.
    """
    import subprocess
    import sys
    
    print("Installing required dependencies for document processing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "PyPDF2", "pandas", "openpyxl"])
    print("Dependencies installed. Please restart your application.")

# Test function
if __name__ == "__main__":
    # If dependencies are missing, prompt to install them
    if not PDF_AVAILABLE or not DOCX_AVAILABLE or not PANDAS_AVAILABLE:
        print("Some document processing capabilities are unavailable.")
        response = input("Do you want to install the required dependencies? (y/n): ")
        if response.lower() == 'y':
            install_dependencies()
            import sys
            sys.exit(0)
    
    # Test processing a file
    import sys
    
    if len(sys.argv) > 1:
        # Process the file provided on command line
        file_path = sys.argv[1]
        print(f"Processing file: {file_path}")
        result = process_file(file_path)
        if result["success"]:
            print(f"Success! Added {result['chunks']} chunks from {result['filename']} ({result['file_type']})")
        else:
            print(f"Failed: {result['error']}")
    else:
        # No command line argument, create a sample text file
        test_file = "example.txt"
        
        if os.path.exists(test_file):
            print(f"Processing test file: {test_file}")
            result = process_file(test_file)
            if result["success"]:
                print(f"Success! Added {result['chunks']} chunks from {result['filename']}")
            else:
                print(f"Failed: {result['error']}")
        else:
            print(f"Test file {test_file} not found. Creating one for testing.")
            
            # Example of creating a test file
            with open(test_file, 'w', encoding='utf-8') as f:
                f.write("""This is a test document.

It contains multiple paragraphs to demonstrate the chunking functionality.

Each paragraph is separated by blank lines.

This document will be used to test the document processing pipeline.

It should be split into chunks based on the specified chunk size and overlap.

The chunks should then be embedded and stored in ChromaDB.

This allows for efficient retrieval based on semantic similarity.

Let's add more text to make sure we create multiple chunks with the default settings.

Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.

Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.

Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.

Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.
""")
            print(f"Created test file {test_file}. Run this script again to process it.")